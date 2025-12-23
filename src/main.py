import sys
import os
from docx import Document
from fpdf import FPDF
import popdf

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QFrame, QLabel, QPushButton, QProgressBar,
    QVBoxLayout, QHBoxLayout, QGridLayout, QFileDialog, QMessageBox, QListWidget,
    QListWidgetItem
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QPixmap, QDragEnterEvent, QDropEvent

# 尝试导入转换库，缺失时提供友好提示
try:
    from pdf2docx import Converter
    import pdfplumber
    import openpyxl
    from PIL import Image
    import fitz  # PyMuPDF

    CONVERSION_ENABLED = True
except ImportError as e:
    CONVERSION_ENABLED = False
    MISSING_MODULE = str(e).split("'")[1]  # 获取缺失的模块名


class PDF(FPDF):
    """自定义PDF类，支持中文显示（解决乱码问题）"""

    def header(self):
        pass  # 可自定义页眉

    def footer(self):
        pass  # 可自定义页脚


class ConversionThread(QThread):
    """转换线程（避免UI卡顿）"""
    progress_update = pyqtSignal(int)
    finished_signal = pyqtSignal(bool, str)
    conversion_type = ""
    input_file = ""
    PDF_output_file = ""
    Word_output_file = ""
    Excel_output_file = ""

    def run(self):
        try:
            if not CONVERSION_ENABLED:
                raise Exception(f"缺少转换依赖库，请先安装：{MISSING_MODULE}")

            if self.conversion_type == "pdf2word":
                self.pdf_to_word()
            elif self.conversion_type == "pdf2excel":
                self.pdf_to_excel()
            elif self.conversion_type == "word2pdf":
                self.word_to_pdf()
            elif self.conversion_type == "excel2pdf":
                self.excel_to_pdf()
            self.finished_signal.emit(True, f"转换完成：\n{self.get_output_path()}")
        except Exception as e:
            self.finished_signal.emit(False, f"转换失败：\n{str(e)}")

    def get_output_path(self):
        """获取当前转换的输出路径"""
        if self.conversion_type == "pdf2word":
            return self.Word_output_file
        elif self.conversion_type == "pdf2excel":
            return self.Excel_output_file
        else:
            return self.PDF_output_file

    def pdf_to_word(self):
        """PDF转 word"""
        if os.path.exists(self.Word_output_file):
            try:
                os.remove(self.Word_output_file)  # 删除旧文件
            except Exception as e:
                raise Exception(f"无法删除旧Word文件：{e}，请关闭该文件后重试")

        cv = Converter(self.input_file)
        pdf_doc = fitz.open(self.input_file)
        total_pages = len(pdf_doc)
        print(f"检测到PDF页数：{total_pages}")

        # 分步转换并更新进度
        for i in range(total_pages):
            cv.convert(self.Word_output_file)
            self.progress_update.emit(int((i + 1) / total_pages * 100))
            print(f"已转换第 {i + 1} 页")  # 调试用：确认逐页执行
        print(f"转换完成")
        cv.close()
        pdf_doc.close()

    def pdf_to_excel(self):
        """PDF转Excel"""
        workbook = openpyxl.Workbook()
        worksheet = workbook.active
        worksheet.title = "PDF内容"

        with pdfplumber.open(self.input_file) as pdf:
            total_pages = len(pdf.pages)
            row = 1
            for i, page in enumerate(pdf.pages):
                try:
                    text = page.extract_text()
                    if text:
                        for line in text.split('\n'):
                            worksheet.cell(row=row, column=1, value=line)
                            row += 1
                    # 更新进度
                    progress = int((i + 1) / total_pages * 100)
                    self.progress_update.emit(progress)
                except Exception as e:
                    self.progress_update.emit(int((i + 1) / total_pages * 100))
                    continue

        workbook.save(self.Excel_output_file)
        self.progress_update.emit(100)

    def word_to_pdf(self):
        """Word转PDF"""
        # 读取Word文档内容
        doc = Document(self.input_file)
        # 初始化PDF对象，设置页面格式和中文字体
        pdf = PDF('P', 'mm', 'A4')  # 纵向、毫米、A4纸张
        pdf.add_page()

        # 字体路径适配（避免找不到字体导致崩溃）
        font_paths = [
            r"C:\Windows\Fonts\simhei.ttf",  # 黑体
            r"C:\Windows\Fonts\msyh.ttf",  # 微软雅黑
            r"C:\Windows\Fonts\simsun.ttc"  # 宋体（备选）
        ]
        font_path = None
        for path in font_paths:
            if os.path.exists(path):
                font_path = path
                break
        if not font_path:
            raise Exception("未找到支持中文的字体文件，请检查系统字体")

        pdf.add_font('SimHei', '', font_path, uni=True)  # 支持中文
        pdf.set_font('SimHei', size=12)  # 设置字体和大小
        line_spacing = 5  # 行间距

        # 遍历Word段落，写入PDF
        total_paragraphs = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                pdf.ln(line_spacing)  # 空行
                continue
            # 自动换行写入中文文本
            pdf.multi_cell(0, 10, txt=para.text, align='L')
            pdf.ln(line_spacing)  # 段落间距
            # 更新进度
            progress = int((i + 1) / total_paragraphs * 100)
            self.progress_update.emit(progress)

        # 保存PDF文件
        pdf.output(self.PDF_output_file)
        self.progress_update.emit(100)


class PDFConverterGUI(QMainWindow):

    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        # 主窗口设置
        self.setWindowTitle("PDF转换器 - 多功能格式转换工具")
        self.setGeometry(100, 100, 1080, 720)
        self.setMinimumSize(720, 480)

        # 中心窗口
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 主布局
        main_layout = QGridLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(10, 10, 10, 10)

        # 配置网格权重
        main_layout.setColumnStretch(0, 1)
        main_layout.setColumnStretch(1, 14)
        main_layout.setRowStretch(0, 1)
        main_layout.setRowStretch(1, 9)

        # 创建组件
        self.create_top_frame(main_layout)
        self.create_left_frame(main_layout)
        self.create_middle_frame(main_layout)

    def create_top_frame(self, parent_layout):
        """顶部标题栏"""
        top_frame = QFrame()
        top_frame.setStyleSheet("background-color: #3c3f41")
        parent_layout.addWidget(top_frame, 0, 0, 1, 2)
        top_layout = QVBoxLayout(top_frame)
        img_text_layout = QHBoxLayout()

        # 加载图片（若不存在则仅显示文字）
        img_label = QLabel(top_frame)
        img_label.setFixedSize(50, 50)
        try:
            img = QPixmap("PDFconverter.ico")
            img = img.scaled(50, 50, Qt.AspectRatioMode.IgnoreAspectRatio, Qt.TransformationMode.SmoothTransformation)
            img_label.setPixmap(img)
        except:
            pass  # 图片不存在时不显示
        img_label.setStyleSheet("border:0.5px solid #ffffff")

        # 顶部文字
        title_label = QLabel("PDF转换器")
        title_font = QFont("微软雅黑", 14, QFont.Weight.Bold)
        title_label.setFont(title_font)

        img_text_layout.addWidget(img_label)
        img_text_layout.addWidget(title_label)

        v_layout = QVBoxLayout()
        v_layout.addSpacing(5)
        h_layout = QHBoxLayout()
        h_layout.addLayout(img_text_layout)
        v_layout.addLayout(h_layout)
        v_layout.addStretch(1)

        top_layout.addLayout(v_layout)

    def create_left_frame(self, parent_layout):
        """左侧功能栏"""
        left_frame = QFrame()
        left_frame.setStyleSheet("background-color: #3c3f41")
        parent_layout.addWidget(left_frame, 1, 0)

        left_layout = QVBoxLayout(left_frame)
        left_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        left_layout.setContentsMargins(10, 20, 10, 10)
        left_layout.setSpacing(10)

        # 功能标题
        func_label = QLabel("功能选择")
        func_font = QFont("微软雅黑", 14, QFont.Weight.Bold)
        func_label.setFont(func_font)
        func_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        left_layout.addWidget(func_label)
        left_layout.addSpacing(10)

        # 按钮样式
        btn_style = """
            QPushButton {
                font-family: 微软雅黑;
                font-size: 12px;
                padding: 8px;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                color: #000;
            }
            QPushButton:pressed {
                background-color: #d0d0d0;
            }
            QPushButton:disabled {
                background-color: #cccccc;
                color: #666666;
            }
        """

        # 功能按钮（依赖缺失时禁用）
        self.pdf2word_btn = QPushButton("PDF转Word")
        self.pdf2word_btn.setStyleSheet(btn_style)
        self.pdf2word_btn.clicked.connect(lambda: self.switch_to_select_func("pdf2word"))
        self.pdf2word_btn.setEnabled(CONVERSION_ENABLED)
        left_layout.addWidget(self.pdf2word_btn)

        self.pdf2excel_btn = QPushButton("PDF转Excel")
        self.pdf2excel_btn.setStyleSheet(btn_style)
        self.pdf2excel_btn.clicked.connect(lambda: self.switch_to_select_func("pdf2excel"))
        self.pdf2excel_btn.setEnabled(CONVERSION_ENABLED)
        left_layout.addWidget(self.pdf2excel_btn)

        self.word2pdf_btn = QPushButton("Word转PDF")
        self.word2pdf_btn.setStyleSheet(btn_style)
        self.word2pdf_btn.clicked.connect(lambda: self.switch_to_select_func("word2pdf"))
        self.word2pdf_btn.setEnabled(CONVERSION_ENABLED)
        left_layout.addWidget(self.word2pdf_btn)

        self.excel2pdf_btn = QPushButton("Excel转PDF")
        self.excel2pdf_btn.setStyleSheet(btn_style)
        self.excel2pdf_btn.clicked.connect(lambda: self.switch_to_select_func("excel2pdf"))
        self.excel2pdf_btn.setEnabled(CONVERSION_ENABLED and sys.platform.startswith('win'))
        left_layout.addWidget(self.excel2pdf_btn)

        left_layout.addStretch()

    def create_middle_frame(self, parent_layout):
        """中间主界面"""
        middle_frame = QFrame()
        middle_frame.setStyleSheet("background-color: #2b2d30")
        parent_layout.addWidget(middle_frame, 1, 1)

        middle_layout = QVBoxLayout(middle_frame)
        middle_layout.setContentsMargins(20, 20, 20, 20)
        middle_layout.setSpacing(20)

        # 最近文档标题
        recent_label = QLabel("最近文档")
        recent_font = QFont("微软雅黑", 14, QFont.Weight.Bold)
        recent_label.setFont(recent_font)
        recent_label.setStyleSheet("background-color: #2b2d30; color: white;")
        middle_layout.addWidget(recent_label)

        # 空状态提示
        empty_label = QLabel("请从左侧选择转换功能开始使用")
        empty_label.setStyleSheet("color: #aaa; font-size: 16px;")
        empty_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        middle_layout.addStretch()
        middle_layout.addWidget(empty_label)
        middle_layout.addStretch()

    def switch_to_select_func(self, conversion_type):
        """跳转到转换功能窗口"""
        if hasattr(self, 'selectfunc') and self.selectfunc.isVisible():
            self.selectfunc.activateWindow()
            self.selectfunc.raise_()
        else:
            self.selectfunc = SelectFunc(conversion_type, self)
            self.selectfunc.show()


class SelectFunc(QMainWindow):
    """转换功能窗口"""

    def __init__(self, conversion_type, main_windows):
        super().__init__()
        self.conversion_type = conversion_type
        self.main_window = main_windows
        self.file_paths = []  # 存储多选文件路径
        self.drag_pos = None  # 窗口拖动位置
        self.init_ui()
        self.move_to_main_window_center()

        # 初始化转换线程
        if CONVERSION_ENABLED:
            self.conversion_thread = ConversionThread()
            self.conversion_thread.progress_update.connect(self.update_progress)
            self.conversion_thread.finished_signal.connect(self.conversion_finished)
        else:
            QMessageBox.warning(
                self,
                "功能受限",
                f"缺少必要的转换库：{MISSING_MODULE}\n\n请执行以下命令安装：\n"
                f"pip install {MISSING_MODULE} -i https://pypi.tuna.tsinghua.edu.cn/simple"
            )

    def move_to_main_window_center(self):
        """将窗口移动到主窗口中心"""
        main_geo = self.main_window.geometry()
        self_geo = self.geometry()
        center_x = main_geo.x() + (main_geo.width() - self_geo.width()) // 2
        center_y = main_geo.y() + (main_geo.height() - self_geo.height()) // 2
        self.move(center_x, center_y)

    def init_ui(self):
        """初始化UI"""
        self.setWindowTitle(self.get_conversion_title())
        self.setGeometry(0, 0, 800, 600)

        # 中心部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        layout = QVBoxLayout(central_widget)

        # 子网格布局
        second_layout = QGridLayout()
        second_layout.setSpacing(10)
        second_layout.setContentsMargins(10, 10, 10, 10)
        second_layout.setRowStretch(0, 1)
        second_layout.setRowStretch(1, 9)
        second_layout.setRowStretch(2, 3)
        layout.addLayout(second_layout)

        self.create_top_frame(second_layout)
        self.create_middle_frame(second_layout)
        self.create_bottom_frame(second_layout)

    def create_top_frame(self, parent_layout):
        """顶部标题框架"""
        top_frame = QFrame()
        top_frame.setStyleSheet("background-color: #3c3f41")
        parent_layout.addWidget(top_frame, 0, 0, 1, 2)
        top_layout = QVBoxLayout(top_frame)

        title_text = f"✨ {self.get_conversion_title()}"
        title = QLabel(title_text)
        title.setStyleSheet("font-size: 20px; color: #2E86AB;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        top_layout.addWidget(title)

    def create_middle_frame(self, parent_layout):
        """中间文件列表框架"""
        middle_frame = QFrame()
        middle_frame.setStyleSheet("background-color: #3c3f41")
        parent_layout.addWidget(middle_frame, 1, 0, 1, 2)

        middle_layout = QVBoxLayout(middle_frame)
        middle_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        middle_layout.setContentsMargins(10, 20, 10, 10)
        middle_layout.setSpacing(10)

        # 顶部工具条
        top_tool_layout = QHBoxLayout()
        list_tip_label = QLabel("已选择的文件：")
        list_tip_label.setStyleSheet("font-size: 14px; color: #2E86AB; font-weight: bold;")
        top_tool_layout.addWidget(list_tip_label)
        top_tool_layout.addStretch()

        # 删除按钮
        delete_btn = QPushButton("删除选中文件")
        delete_btn.clicked.connect(self.delete_selected_file)
        delete_btn.setStyleSheet("""
            QPushButton {
                padding: 8px 16px;
                font-size: 14px;
                background-color: #f44336;
                color: white;
                border: none;
                border-radius: 6px;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        top_tool_layout.addWidget(delete_btn)
        middle_layout.addLayout(top_tool_layout)

        # 文件列表
        self.file_list_widget = QListWidget()
        self.file_list_widget.setMinimumHeight(200)
        self.file_list_widget.setSelectionMode(QListWidget.SelectionMode.MultiSelection)

        # 开启拖放接受功能
        self.file_list_widget.setAcceptDrops(True)
        # 隐藏默认拖放指示器（如需自定义样式）
        self.file_list_widget.setDropIndicatorShown(True)

        # 补充QListWidget样式，与深色主题统一
        self.file_list_widget.setStyleSheet("""
                   QListWidget {
                       background-color: #2b2b2b;
                       color: #ffffff;
                       font-size: 13px;
                       border: none;
                       border-radius: 6px;
                       padding: 5px;
                   }
                   QListWidget::item:selected {
                       background-color: #2E86AB;
                       color: white;
                   }
                   QListWidget::item:hover {
                       background-color: #4a4d4f;
                   }
               """)
        middle_layout.addWidget(self.file_list_widget)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid grey;
                border-radius: 5px;
                text-align: center;
                height: 20px;
            }
            QProgressBar::chunk {
                background-color: #4CAF50;
                width: 10px;
            }
        """)
        self.progress_bar.setValue(0)
        middle_layout.addWidget(self.progress_bar)

        # 判断拖入数据是否合法

    def dragEnterEvent(self, event: QDragEnterEvent):
        # 判断是否为文件路径数据
        if event.mimeData().hasUrls():
            event.acceptProposedAction()  # 允许拖放
        else:
            event.ignore()  # 忽略无效拖放

        # 支持拖放过程中移动

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            event.ignore()

        # 处理拖放落地逻辑，添加文件到列表

    def dropEvent(self, event: QDropEvent):
        # 获取拖入的所有数据（URL格式）
        mime_data = event.mimeData()
        if not mime_data.hasUrls():
            event.ignore()
            return

        # 遍历所有URL，转换为本地文件路径
        for url in mime_data.urls():
            # 关键：将QUrl转换为本地文件路径（解决中文路径乱码问题）
            file_path = url.toLocalFile()
            # 过滤：只添加实际存在的文件（排除文件夹）
            if os.path.isfile(file_path):
                # 添加文件路径（也可只添加文件名：os.path.basename(file_path)）
                self.file_list_widget.addItem(file_path)

        event.acceptProposedAction()

    def create_bottom_frame(self, parent_layout):
        """底部按钮框架"""
        bottom_frame = QFrame()
        bottom_frame.setStyleSheet("background-color: #3c3f41")
        parent_layout.addWidget(bottom_frame, 2, 0, 1, 2)
        bottom_layout = QHBoxLayout(bottom_frame)
        bottom_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        bottom_layout.setContentsMargins(50, 20, 50, 10)
        bottom_layout.setSpacing(30)

        # 选择文件按钮
        select_btn = QPushButton("选择文件")
        select_btn.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                font-size: 16px;
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        select_btn.clicked.connect(lambda: self.select_file(self.conversion_type))

        # 转换按钮
        converter_btn = QPushButton("开始转换")
        converter_btn.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                font-size: 16px;
                background-color: #4CAF50;
                color: white;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        converter_btn.clicked.connect(lambda: self.converter_func(self.conversion_type))

        # 返回按钮
        back_btn = QPushButton("返回主窗口")
        back_btn.setStyleSheet("""
            QPushButton {
                padding: 10px 20px;
                font-size: 16px;
                background-color: #2196F3;
                color: white;
                border: none;
                border-radius: 5px;
            }
            QPushButton:hover {
                background-color: #0b7dda;
            }
        """)
        back_btn.clicked.connect(self.back_to_main)

        bottom_layout.addWidget(select_btn)
        bottom_layout.addWidget(converter_btn)
        bottom_layout.addWidget(back_btn)

    def get_conversion_title(self):
        """获取转换类型标题"""
        title_map = {
            "pdf2word": "PDF转Word 转换界面",
            "pdf2excel": "PDF转Excel 转换界面",
            "word2pdf": "Word转PDF 转换界面",
            "excel2pdf": "Excel转PDF 转换界面"
        }
        return title_map.get(self.conversion_type, "PDF转换界面")

    def mousePressEvent(self, event):
        """窗口拖动：鼠标按下"""
        if event.button() == Qt.MouseButton.LeftButton:
            self.drag_pos = event.globalPosition().toPoint() - self.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        """窗口拖动：鼠标移动"""
        if event.buttons() == Qt.MouseButton.LeftButton and self.drag_pos is not None:
            self.move(event.globalPosition().toPoint() - self.drag_pos)
            event.accept()

    def select_file(self, conversion_type):
        """选择文件"""
        if conversion_type in ["pdf2word", "pdf2excel"]:
            # PDF转其他格式
            file_paths, _ = QFileDialog.getOpenFileNames(
                self, "选择PDF文件（可多选）", "", "PDF文件 (*.pdf);;所有文件 (*.*)"
            )
        elif conversion_type == "word2pdf":
            # Word转PDF
            file_paths, _ = QFileDialog.getOpenFileNames(
                self, "选择Word文件（可多选）", "", "Word文件 (*.docx *.doc);;所有文件 (*.*)"
            )
        elif conversion_type == "excel2pdf":
            # Excel转PDF
            file_paths, _ = QFileDialog.getOpenFileNames(
                self, "选择Excel文件（可多选）", "", "Excel文件 (*.xlsx *.xls);;所有文件 (*.*)"
            )
        else:
            return

        # 添加文件到列表
        for path in file_paths:
            if path not in self.file_paths:
                self.file_paths.append(path)
                self.file_list_widget.addItem(QListWidgetItem(path))

    def delete_selected_file(self):
        """删除选中文件"""
        selected_items = self.file_list_widget.selectedItems()
        if not selected_items:
            return  # 无选中项时直接返回
        # 反向删除避免正向删除导致索引错乱
        for item in reversed(selected_items):
            row = self.file_list_widget.row(item)
            self.file_list_widget.takeItem(row)

    def back_to_main(self):
        """返回主窗口"""
        self.close()
        self.main_window.activateWindow()

    def update_progress(self, value):
        """更新进度条"""
        self.progress_bar.setValue(value)

    def conversion_finished(self, success, message):
        """转换完成回调"""
        if success:
            QMessageBox.information(self, "成功", message)
        else:
            QMessageBox.critical(self, "失败", message)
        self.progress_bar.setValue(0)

    def converter_func(self, conversion_type):
        """转换功能入口"""
        if not self.file_paths:
            QMessageBox.warning(self, "警告", "请先选择文件")
            return

        # 确保线程未运行
        if hasattr(self, 'conversion_thread') and self.conversion_thread.isRunning():
            QMessageBox.information(self, "提示", "转换正在进行中，请稍后再试")
            return

        # 批量转换文件
        for input_file in self.file_paths:
            # 生成输出路径
            if conversion_type == "pdf2word":
                output_file = os.path.splitext(input_file)[0] + ".docx"
            elif conversion_type == "pdf2excel":
                output_file = os.path.splitext(input_file)[0] + ".xlsx"
            elif conversion_type == "word2pdf":
                output_file = os.path.splitext(input_file)[0] + ".pdf"

            # 校验输出路径
            if not output_file or os.path.isdir(output_file):
                QMessageBox.warning(self, "错误", f"无效的输出路径：{output_file}")
                continue

            # 配置线程参数
            self.conversion_thread = ConversionThread()
            self.conversion_thread.conversion_type = conversion_type
            self.conversion_thread.input_file = input_file
            self.conversion_thread.progress_update.connect(self.update_progress)
            self.conversion_thread.finished_signal.connect(self.conversion_finished)

            # 设置输出路径
            if conversion_type == "pdf2word":
                self.conversion_thread.Word_output_file = output_file
            elif conversion_type == "pdf2excel":
                self.conversion_thread.Excel_output_file = output_file
            elif conversion_type == "word2pdf":
                self.conversion_thread.PDF_output_file = output_file

            # 启动线程
            self.conversion_thread.start()
            # 等待当前文件转换完成（避免多线程冲突）
            self.conversion_thread.wait()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PDFConverterGUI()
    window.show()
    sys.exit(app.exec())
